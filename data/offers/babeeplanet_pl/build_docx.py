#!/usr/bin/env python3
"""Auto-generated markdown-to-docx converter dla babeeplanet_pl."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SLUG = "babeeplanet_pl"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OFFER_MD  = os.path.join(BASE_DIR, "offer.md")
OUTPUT    = os.path.join(BASE_DIR, f"offer_{SLUG}.docx")

# Brand colors Double Digital
ORANGE   = RGBColor(0xFF, 0x6B, 0x00)
DARK     = RGBColor(0x1A, 0x1A, 0x1A)
MID_GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Cm(2.5)
    s.top_margin  = s.bottom_margin = Cm(2.0)

# Default font
for style_name in ["Normal", "Body Text"]:
    try:
        doc.styles[style_name].font.name = "Calibri"
        doc.styles[style_name].font.size = Pt(11)
    except Exception:
        pass


def add_heading(text, level):
    p = doc.add_heading(text.strip(), level=level)
    if p.runs:
        if level == 1:
            p.runs[0].font.color.rgb = ORANGE
        elif level == 2:
            p.runs[0].font.color.rgb = DARK
        else:
            p.runs[0].font.color.rgb = MID_GREY


def add_rich_para(text, style=None):
    """Adds paragraph with **bold** markdown support."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for idx, part in enumerate(parts):
        run = p.add_run(part)
        if idx % 2 == 1:
            run.bold = True
    return p


def flush_table(buf):
    rows = [l for l in buf if l.strip().startswith("|") and not re.match(r"^\|[-| ]+\|$", l.strip())]
    if not rows:
        return
    cols = [c.strip() for c in rows[0].split("|") if c.strip()]
    n_cols = len(cols)
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = "Table Grid"
    for r_idx, row_str in enumerate(rows):
        cells = [c.strip() for c in row_str.split("|") if c.strip()]
        for c_idx, cell_txt in enumerate(cells[:n_cols]):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = cell_txt
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


with open(OFFER_MD, encoding="utf-8") as f:
    all_lines = f.readlines()

table_buf = []
for raw in all_lines:
    line = raw.rstrip("\n")

    if line.strip().startswith("|"):
        table_buf.append(line)
        continue
    elif table_buf:
        flush_table(table_buf)
        table_buf = []

    if   line.startswith("# "):    add_heading(line[2:], 1)
    elif line.startswith("## "):   add_heading(line[3:], 2)
    elif line.startswith("### "):  add_heading(line[4:], 3)
    elif line.startswith("#### "): add_heading(line[5:], 4)
    elif line.startswith("> "):
        add_rich_para(line[2:], style="Quote")
    elif re.match(r"^[-*] ", line):
        add_rich_para(line[2:], style="List Bullet")
    elif line.strip() == "---":
        doc.add_paragraph()
    elif line.strip():
        add_rich_para(line)

if table_buf:
    flush_table(table_buf)

doc.save(OUTPUT)
print(f"Zapisano: {OUTPUT}")
