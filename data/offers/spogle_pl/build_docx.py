"""
Oferta SEO Double Digital — spogle.pl
CRO-first layout: value prop w 5 sek, scannability, objection handling, CTA
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Paleta ────────────────────────────────────────────
C_BLACK  = RGBColor(0x0F, 0x0F, 0x0F)
C_ACCENT = RGBColor(0xFF, 0x6B, 0x00)
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_GREY   = RGBColor(0x4A, 0x4A, 0x4A)
C_LGREY  = RGBColor(0xF7, 0xF7, 0xF7)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN  = RGBColor(0x16, 0xA3, 0x4A)
C_RED    = RGBColor(0xDC, 0x26, 0x26)
C_ORANGE = RGBColor(0xEA, 0x58, 0x0C)
C_YELLOW = RGBColor(0xB4, 0x5B, 0x09)
C_TEAL   = RGBColor(0x06, 0x82, 0x72)
C_LBLUE  = RGBColor(0xEF, 0xF6, 0xFF)
C_LGREEN = RGBColor(0xF0, 0xFD, 0xF4)
C_LRED   = RGBColor(0xFE, 0xF2, 0xF2)
C_SAND   = RGBColor(0xFD, 0xF6, 0xEE)

FONT = "Calibri"

# ── Core helpers ──────────────────────────────────────

def rgb_hex(c): return str(c)

def margins(doc, t=1.8, b=2.0, l=2.4, r=2.4):
    for s in doc.sections:
        s.top_margin=Cm(t); s.bottom_margin=Cm(b)
        s.left_margin=Cm(l); s.right_margin=Cm(r)

def sp(para, before=0, after=5):
    f = para.paragraph_format
    f.space_before=Pt(before); f.space_after=Pt(after)

def r(para, text, bold=False, italic=False, sz=11, color=None):
    rn = para.add_run(text)
    rn.bold=bold; rn.italic=italic
    rn.font.name=FONT; rn.font.size=Pt(sz)
    if color: rn.font.color.rgb=color
    return rn

def shd_para(p_elem, rgb):
    pPr=p_elem.get_or_add_pPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),rgb_hex(rgb))
    pPr.append(s)

def shd_cell(cell, rgb):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),rgb_hex(rgb))
    tcPr.append(s)

def bar_left(para, rgb, sz='20'):
    pPr=para._p.get_or_add_pPr()
    pBdr=OxmlElement('w:pBdr')
    l=OxmlElement('w:left')
    l.set(qn('w:val'),'single'); l.set(qn('w:sz'),sz)
    l.set(qn('w:space'),'6'); l.set(qn('w:color'),rgb_hex(rgb))
    pBdr.append(l); pPr.append(pBdr)

def bar_bottom(para, rgb, sz='10'):
    pPr=para._p.get_or_add_pPr()
    pBdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz)
    b.set(qn('w:space'),'3'); b.set(qn('w:color'),rgb_hex(rgb))
    pBdr.append(b); pPr.append(pBdr)

def gap(doc, h=5):
    p=doc.add_paragraph(); sp(p,before=h,after=0); return p

# ── Typografia ────────────────────────────────────────

def H1(doc, text, color=C_DARK, accent=C_ACCENT):
    p=doc.add_paragraph(); sp(p,before=18,after=5)
    bar_bottom(p, accent, sz='14')
    rn=p.add_run(text); rn.bold=True
    rn.font.name=FONT; rn.font.size=Pt(19); rn.font.color.rgb=color
    return p

def H2(doc, text, color=C_DARK):
    p=doc.add_paragraph(); sp(p,before=12,after=3)
    rn=p.add_run(text); rn.bold=True
    rn.font.name=FONT; rn.font.size=Pt(13); rn.font.color.rgb=color
    return p

def H3(doc, text, color=C_GREY):
    p=doc.add_paragraph(); sp(p,before=8,after=2)
    rn=p.add_run(text); rn.bold=True
    rn.font.name=FONT; rn.font.size=Pt(11); rn.font.color.rgb=color
    return p

def body(doc, text, color=C_GREY, sz=10.5, before=0, after=5, italic=False):
    p=doc.add_paragraph(); sp(p,before=before,after=after)
    p.paragraph_format.left_indent=Cm(0)
    r(p,text,sz=sz,color=color,italic=italic)
    return p

def bul(doc, text, color=C_GREY, bold_part=None, indent=0.7):
    p=doc.add_paragraph(style='List Bullet')
    sp(p,before=1,after=2)
    p.paragraph_format.left_indent=Cm(indent)
    if bold_part:
        r(p,bold_part,bold=True,sz=10.5,color=C_BLACK)
        r(p,text,sz=10.5,color=color)
    else:
        r(p,text,sz=10.5,color=color)
    return p

# ── Komponenty CRO ────────────────────────────────────

def table(doc, headers, rows, widths=None, hbg=C_DARK, fontsize=9.5):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''
        shd_cell(c,hbg); p=c.paragraphs[0]
        rn=p.add_run(h); rn.bold=True
        rn.font.name=FONT; rn.font.size=Pt(fontsize); rn.font.color.rgb=C_WHITE
    for ri,row in enumerate(rows):
        bg=C_LGREY if ri%2==0 else C_WHITE
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''
            shd_cell(c,bg); p=c.paragraphs[0]
            rn=p.add_run(str(val))
            rn.font.name=FONT; rn.font.size=Pt(fontsize); rn.font.color.rgb=C_GREY
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    gap(doc,4); return t

def callout(doc, text, bg=C_LGREY, bar=C_ACCENT, italic=False, sz=10.5, bold=False):
    p=doc.add_paragraph(); sp(p,before=5,after=5)
    p.paragraph_format.left_indent=Cm(0.5)
    p.paragraph_format.right_indent=Cm(0.5)
    shd_para(p._p,bg); bar_left(p,bar,sz='28')
    r(p,text,italic=italic,sz=sz,color=C_BLACK,bold=bold)
    return p

def big_stat(doc, items):
    """Karty ze statystykami — ikona + liczba + opis"""
    t=doc.add_table(rows=1,cols=len(items))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,(icon,num,label,bg) in enumerate(items):
        c=t.rows[0].cells[i]; c.text=''
        shd_cell(c,bg)
        pi=c.add_paragraph(); sp(pi,before=6,after=0)
        pi.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r(pi,icon,sz=20)
        pn=c.add_paragraph(); sp(pn,before=3,after=0)
        pn.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r(pn,num,bold=True,sz=20,color=C_DARK)
        pl=c.add_paragraph(); sp(pl,before=0,after=6)
        pl.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r(pl,label,sz=9,color=C_GREY)
    gap(doc,6); return t

def hero_number(doc, number, unit, context, bg=C_DARK, num_color=C_ACCENT):
    """Duża liczba z kontekstem — CRO pull quote"""
    t=doc.add_table(rows=1,cols=2)
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    left=t.rows[0].cells[0]; right=t.rows[0].cells[1]
    shd_cell(left,bg); shd_cell(right,C_LGREY)
    left.width=Cm(5); right.width=Cm(12)
    pn=left.add_paragraph(); sp(pn,before=12,after=0)
    pn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pn,number,bold=True,sz=36,color=num_color)
    pu=left.add_paragraph(); sp(pu,before=0,after=12)
    pu.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pu,unit,sz=10,color=C_WHITE)
    pc=right.add_paragraph(); sp(pc,before=6,after=6)
    pc.paragraph_format.left_indent=Cm(0.5)
    r(pc,context,sz=11,color=C_GREY)
    gap(doc,5); return t

def section_badge(doc, icon, label, bg, text_color=C_WHITE):
    """Kolorowy label sekcji"""
    p=doc.add_paragraph(); sp(p,before=12,after=0)
    shd_para(p._p,bg)
    p.paragraph_format.left_indent=Cm(0.4)
    r(p,f"  {icon}  {label}  ",bold=True,sz=11,color=text_color)
    return p

def problem_card(doc, num, title, what, effect, scale, bg_left=C_RED):
    """Karta problemu — what/effect/scale"""
    t=doc.add_table(rows=1,cols=2)
    t.style='Table Grid'
    left=t.rows[0].cells[0]; right=t.rows[0].cells[1]
    shd_cell(left,bg_left); shd_cell(right,C_LGREY)
    left.width=Cm(2); right.width=Cm(15)
    pn=left.add_paragraph(); sp(pn,before=10,after=0)
    pn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pn,num,bold=True,sz=28,color=C_WHITE)
    pt=left.add_paragraph(); sp(pt,before=0,after=10)
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pt,"PROBLEM",sz=7,color=C_WHITE)
    # right
    pt2=right.add_paragraph(); sp(pt2,before=6,after=2)
    pt2.paragraph_format.left_indent=Cm(0.4)
    r(pt2,title,bold=True,sz=12,color=C_DARK)
    pw=right.add_paragraph(); sp(pw,before=2,after=2)
    pw.paragraph_format.left_indent=Cm(0.4)
    r(pw,"Co się dzieje: ",bold=True,sz=10,color=C_BLACK)
    r(pw,what,sz=10,color=C_GREY)
    pe=right.add_paragraph(); sp(pe,before=2,after=2)
    pe.paragraph_format.left_indent=Cm(0.4)
    r(pe,"Skutek: ",bold=True,sz=10,color=C_BLACK)
    r(pe,effect,sz=10,color=C_GREY)
    ps=right.add_paragraph(); sp(ps,before=2,after=6)
    ps.paragraph_format.left_indent=Cm(0.4)
    r(ps,"Skala: ",bold=True,sz=10,color=C_BLACK)
    r(ps,scale,sz=10,color=bg_left)
    gap(doc,4); return t

def opportunity_card(doc, letter, title, timing, desc, potential):
    t=doc.add_table(rows=1,cols=2)
    t.style='Table Grid'
    left=t.rows[0].cells[0]; right=t.rows[0].cells[1]
    shd_cell(left,C_ACCENT); shd_cell(right,C_LGREY)
    left.width=Cm(2); right.width=Cm(15)
    pl=left.add_paragraph(); sp(pl,before=10,after=0)
    pl.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pl,letter,bold=True,sz=26,color=C_WHITE)
    pt2=left.add_paragraph(); sp(pt2,before=0,after=10)
    pt2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pt2,timing,sz=7,color=C_WHITE)
    pt=right.add_paragraph(); sp(pt,before=6,after=2)
    pt.paragraph_format.left_indent=Cm(0.4)
    r(pt,title,bold=True,sz=12,color=C_DARK)
    pd=right.add_paragraph(); sp(pd,before=0,after=2)
    pd.paragraph_format.left_indent=Cm(0.4)
    r(pd,desc,sz=10,color=C_GREY)
    pp=right.add_paragraph(); sp(pp,before=2,after=6)
    pp.paragraph_format.left_indent=Cm(0.4)
    r(pp,"Potencjał: ",bold=True,sz=10,color=C_GREEN)
    r(pp,potential,sz=10,color=C_GREEN)
    gap(doc,4); return t

def objection_row(doc, question, answer):
    t=doc.add_table(rows=1,cols=2)
    t.style='Table Grid'
    left=t.rows[0].cells[0]; right=t.rows[0].cells[1]
    shd_cell(left,C_SAND); shd_cell(right,C_WHITE)
    left.width=Cm(6); right.width=Cm(11)
    pq=left.add_paragraph(); sp(pq,before=6,after=6)
    pq.paragraph_format.left_indent=Cm(0.3)
    r(pq,"❓ ",sz=12); r(pq,question,bold=True,sz=10.5,color=C_DARK)
    pa=right.add_paragraph(); sp(pa,before=6,after=6)
    pa.paragraph_format.left_indent=Cm(0.4)
    r(pa,answer,sz=10.5,color=C_GREY)
    gap(doc,2); return t

def action_row(doc, priority_label, timing, bg, title, desc):
    t=doc.add_table(rows=1,cols=2)
    t.style='Table Grid'
    left=t.rows[0].cells[0]; right=t.rows[0].cells[1]
    shd_cell(left,bg); shd_cell(right,C_LGREY)
    left.width=Cm(2.2); right.width=Cm(14.8)
    pl=left.add_paragraph(); sp(pl,before=6,after=0)
    pl.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r(pl,priority_label,bold=True,sz=14,color=C_WHITE)
    pa=right.add_paragraph(); sp(pa,before=6,after=2)
    pa.paragraph_format.left_indent=Cm(0.4)
    r(pa,title,bold=True,sz=11,color=C_DARK)
    pd=right.add_paragraph(); sp(pd,before=0,after=6)
    pd.paragraph_format.left_indent=Cm(0.4)
    r(pd,desc,sz=10,color=C_GREY)
    gap(doc,3); return t

# ══════════════════════════════════════════════════════
# BUDOWANIE DOKUMENTU
# ══════════════════════════════════════════════════════

doc=Document()
margins(doc)

# ─────────────────────────────────────────────────────
# COVER — CRO rule: value prop w 5 sekund, nie logo
# Klient musi zrozumieć "co zyskam" zanim przeczyta cokolwiek
# ─────────────────────────────────────────────────────

# Górny pasek marki
p=doc.add_paragraph(); sp(p,before=0,after=0)
shd_para(p._p,C_DARK)
p.paragraph_format.left_indent=Cm(0.4)
r(p,"  DOUBLE DIGITAL  ·  ANALIZA SEO  ·  2026",sz=8.5,color=RGBColor(0x99,0x99,0x99))

gap(doc,20)

# Headline CRO: outcome-focused, nie "raport SEO"
p=doc.add_paragraph(); sp(p,before=0,after=4)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"Masz ruch.\nNie masz klientów.",bold=True,sz=32,color=C_DARK)

p=doc.add_paragraph(); sp(p,before=4,after=4)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"Oto dlaczego — i jak to naprawić.",italic=True,sz=14,color=C_ACCENT)

gap(doc,10)

# Subheadline z konkretem — CRO: specificity
p=doc.add_paragraph(); sp(p,before=0,after=2)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"spogle.pl",bold=True,sz=22,color=C_DARK)

p=doc.add_paragraph(); sp(p,before=0,after=20)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"57 fraz gotowych do pierwszej strony Google — bez pisania nowych treści",
  sz=11,color=C_GREY,italic=True)

# 3 najważniejsze odkrycia — CRO: scannable above-the-fold summary
big_stat(doc,[
    ("📉","–36%","fraz w Top 10\nw ciągu roku",C_LRED),
    ("⚡","57","fraz na poz. 11–20\ngotowych do awansu",C_LGREEN),
    ("💰","4 755 zł","wartość ruchu org.\n(ekwiwalent Google Ads)",C_LGREY),
])

gap(doc,16)
p=doc.add_paragraph(); sp(p,before=0,after=2)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"Przygotował: Double Digital — Agencja Performance Marketingu",sz=9,color=C_GREY)
p=doc.add_paragraph(); sp(p,before=0,after=0)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"27 lutego 2026  |  Dane: Senuto, NodeHub SERP, analiza treści",sz=8,color=RGBColor(0xAA,0xAA,0xAA),italic=True)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# STRONA 2 — "3 KLUCZOWE ODKRYCIA"
# CRO rule: executive scan page — klient ma 30 sekund, musi wynieść 3 rzeczy
# ─────────────────────────────────────────────────────

H1(doc,"Trzy rzeczy, które musisz wiedzieć")
body(doc,
    "Zanim przejdziemy do szczegółów — oto trzy odkrycia, które kształtują całą "
    "naszą rekomendację. Każde ma bezpośredni wpływ na Twój biznes.",
    before=4,after=12
)

# Odkrycie 1
callout(doc,
    "1 / 3    Twój ruch pochodzi z artykułów, nie z oferty",
    bg=C_LRED, bar=C_RED, bold=True, sz=11
)
p=doc.add_paragraph(); sp(p,before=4,after=2)
p.paragraph_format.left_indent=Cm(0.4)
r(p,"76% widoczności generują 3 artykuły poradnikowe.",bold=True,sz=11,color=C_DARK)
p2=doc.add_paragraph(); sp(p2,before=0,after=10)
p2.paragraph_format.left_indent=Cm(0.4)
r(p2," Ludzie trafiają na blog, czytają, wychodzą. "
    "Twoje strony dmuchańców i bumper balla generują mniej ruchu "
    "niż artykuł o kwarantannie COVID z 2020 roku.",sz=10.5,color=C_GREY)

# Odkrycie 2
callout(doc,
    "2 / 3    Jesteś tuż za pierwszą stroną Google na 57 frazach",
    bg=C_LGREEN, bar=C_GREEN, bold=True, sz=11
)
p=doc.add_paragraph(); sp(p,before=4,after=2)
p.paragraph_format.left_indent=Cm(0.4)
r(p,"To największa szansa, którą identyfikujemy.",bold=True,sz=11,color=C_DARK)
p2=doc.add_paragraph(); sp(p2,before=0,after=10)
p2.paragraph_format.left_indent=Cm(0.4)
r(p2," 57 fraz siedzi na pozycjach 11–20 (tuż za pierwszą stroną) "
    "i ma niską konkurencję. Optymalizacja istniejących stron — "
    "bez pisania nowych artykułów — może je przesunąć do Top 10 w 60–90 dni.",
    sz=10.5,color=C_GREY)

# Odkrycie 3
callout(doc,
    "3 / 3    Twoje podstrony lokalne walczą między sobą",
    bg=C_SAND, bar=C_ORANGE, bold=True, sz=11
)
p=doc.add_paragraph(); sp(p,before=4,after=2)
p.paragraph_format.left_indent=Cm(0.4)
r(p,"59 fraz — problem kanibalizacji.",bold=True,sz=11,color=C_DARK)
p2=doc.add_paragraph(); sp(p2,before=0,after=10)
p2.paragraph_format.left_indent=Cm(0.4)
r(p2," 6 różnych stron (Kraków, Warszawa, Łódź...) walczy o "
    "to samo słowo kluczowe. Google się gubi — żadna nie trafia na pierwszą stronę. "
    "Naprawa to 1–2 tygodnie pracy technicznej, efekt w 60–90 dni.",
    sz=10.5,color=C_GREY)

gap(doc,6)

# CRO: Pull quote z największą liczbą — zatrzymuje oko przy skanowaniu
hero_number(doc,
    "4,4×",
    "więcej widoczności\nniż Twój największy rywal",
    "Spogle.pl (3 677 pkt) vs megaland.pl (844 pkt) — "
    "jesteś liderem swojej niszy. To rzadkość i realny asset, "
    "który warto chronić i rozbudowywać, zanim megaland.pl (+58,6% r/r) "
    "zacznie odrabiać dystans.",
    bg=C_DARK, num_color=C_ACCENT
)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# CZĘŚĆ 1: GDZIE JESTEŚ TERAZ
# CRO rule: najpierw pokaż "rozumiem Twoją sytuację" — empatia przed pitch
# ─────────────────────────────────────────────────────

H1(doc,"1. Gdzie teraz jesteś — pełny obraz")

body(doc,
    "Zanim zaproponujemy cokolwiek — pokaż Ci dokładnie jak wygląda Twoja widoczność "
    "w Google. Liczby z wyjaśnieniem co konkretnie oznaczają dla Twojego biznesu.",
    before=4,after=10
)

H2(doc,"Skąd pochodzi Twój ruch organiczny")

body(doc,
    "Oto 6 stron spogle.pl, które w tej chwili odpowiadają za największą "
    "część Twojej widoczności. Liczba po lewej to frazy rankujące, po prawej — "
    "punkty widoczności Senuto (im więcej, tym więcej kliknięć).",
    after=8
)

# Tabela stron — wizualnie, nie suchare dane
table(doc,
    headers=["Strona","Frazy","Widoczność","Typ","Uwaga"],
    rows=[
        ["/7-zabaw-integracyjnych-dla-doroslych/","129","1 360 pkt","Blog/Poradnik","Twój #1 — artykuł informacyjny"],
        ["/zabawy-integracyje-dla-pracownikow/","105","746 pkt","Blog/Poradnik","Artykuł integracyjny"],
        ["/8-nietypowych-gier-domowych-na-kwarantanne/","109","690 pkt","Blog 2020","⚠ COVID-artykuł — ruch bez wartości"],
        ["/hotel-z-atrakcjami-dla-dzieci-blisko-warszawy/","12","362 pkt","Blog/Lista","Ruch lokalny, sezonowy"],
        ["/uslugi/bumper-ball-wynajem/","10","179 pkt","Strona usługi","Pierwsza strona usługi — na 5. miejscu"],
        ["/uslugi/dmuchance-warszawa/","6","158 pkt","Strona usługi","Dmuchańce Warszawa"],
    ],
    widths=[6.5,2.5,3,3.5,4.5],
)

callout(doc,
    "📌  Co to znaczy w praktyce: Twoje 3 najlepsze strony to artykuły, nie oferty. "
    "Ludzie trafiają, czytają — i wychodzą bez zamówienia. "
    "Strony usług (dmuchańce, bumper ball) generują razem mniej widoczności "
    "niż jeden artykuł o kwarantannie z 2020 roku.",
    bg=C_LBLUE, bar=C_TEAL
)

gap(doc,8)
H2(doc,"Jak zmieniała się Twoja widoczność w ciągu roku")

table(doc,
    headers=["Kiedy","Frazy w Top 10","Frazy w Top 3","Co się stało"],
    rows=[
        ["Luty 2025","815","297","Twój szczyt w ostatnim roku"],
        ["Maj 2025","515","214","⚠ Core Update Google — gwałtowny spadek"],
        ["Teraz (luty 2026)","523","207","Częściowe odbicie, ale wciąż –36% r/r"],
    ],
    widths=[4,3.5,3.5,7],
)

body(doc,
    "Rok temu miałeś o 300 fraz więcej w pierwszej dziesiątce Google. "
    "Aktualizacja algorytmu z maja 2025 uderzyła w strony bez wyraźnych sygnałów "
    "ekspertyzy (E-E-A-T) — to prawdopodobnie główna przyczyna spadku. "
    "Dobra wiadomość: odrabiacie straty, ale potrzeba działań, żeby nie spaść ponownie.",
    before=2,after=10
)

H2(doc,"Ile jest warta Twoja organiczna widoczność?")
body(doc,"Gdybyś chciał kupić taki sam ruch przez Google Ads (reklamy płatne):",after=5)

# CRO: duże, skanowalne liczby — wartość pieniężna przemawia do niespecjalistów
p=doc.add_paragraph(); sp(p,before=2,after=2)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
shd_para(p._p,C_DARK)
r(p,"    4 755 zł / miesiąc    ",bold=True,sz=26,color=C_ACCENT)

body(doc,
    "Tyle wynosi szacunkowy koszt zakupu tego samego ruchu przez Google Ads. "
    "SEO pozwala Ci go zdobywać bezpłatnie — zamiast płacić za każde kliknięcie. "
    "Naszym celem jest zwiększenie tej wartości, nie jej utrzymanie.",
    before=6,after=5
)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# CZĘŚĆ 2: CO CIĘ HAMUJE
# CRO rule: problem agitation — pokaż ból, zanim zaproponujesz lekarstwo
# Każdy problem: co → dlaczego → skutek biznesowy (nie SEO-jargon)
# ─────────────────────────────────────────────────────

H1(doc,"2. Co Cię hamuje?")
body(doc,
    "Zidentyfikowaliśmy 3 główne blokery. Każdy opisujemy w języku biznesowym — "
    "co się dzieje, dlaczego, i co traci Twoja firma.",
    before=4,after=10
)

problem_card(doc,"1",
    "Twoje własne strony konkurują ze sobą w Google",
    what=(
        "Masz 6 stron dla różnych miast (Kraków, Warszawa, Łódź...) "
        "które walczą o identyczne frazy jak \"wynajem dmuchańców na urodziny\". "
        "Google widzi 6 kandydatów do jednej pozycji — i żadnej nie promuje."
    ),
    effect=(
        "Fraza \"wynajem dmuchańców na urodziny\" (320 wyszukiwań/mies.) "
        "ląduje na pozycji 16. — poza pierwszą stroną. Klienci szukają, "
        "ale nie trafiają na Twój serwis — trafiają do konkurentów."
    ),
    scale="59 fraz dotkniętych tym problemem — klasyfikacja Senuto: HIGH severity",
    bg_left=C_RED
)

problem_card(doc,"2",
    "Blog przyciąga czytaczy, nie klientów",
    what=(
        "76% Twojego ruchu z Google pochodzi z artykułów poradnikowych "
        "(zabawy integracyjne, gry domowe). To wartościowy ruch, ale z minimalnym "
        "zamiarem zakupu. Brakuje strony cennikowej, FAQ i ścieżki do zamówienia."
    ),
    effect=(
        "Ktoś czyta artykuł o grach integracyjnych, nie widzi oferty wynajmu dmuchańca, "
        "wychodzi i zamawia u konkurencji. "
        "Twoje strony usług generują mniej ruchu niż artykuł o COVID z 2020 roku."
    ),
    scale="Brak strony cennikowej i FAQ — Google PAA pyta \"ile kosztuje wynajem dmuchańca?\" — brak odpowiedzi na spogle.pl",
    bg_left=C_ORANGE
)

problem_card(doc,"3",
    "Google nie wie, że jesteś ekspertem",
    what=(
        "Strona główna nie mówi Google'owi: \"za nami 10 lat i 1000+ eventów\". "
        "Brak strony /o-nas (aktualnie 404), brak autora w artykułach, "
        "brak dat aktualizacji i zdjęć z realizacji."
    ),
    effect=(
        "Core Update Google z maja 2025 uderzył w strony bez sygnałów ekspertyzy. "
        "Straciłeś 36% fraz Top 10 w 3 miesiące. "
        "Bez naprawy E-E-A-T kolejna aktualizacja może uderzyć ponownie."
    ),
    scale="Brak /o-nas → 404, 0 artykułów z autorem i datą, 0 case studies na stronie",
    bg_left=C_YELLOW
)

gap(doc,6)
callout(doc,
    "✅  Dobra wiadomość: żaden z tych problemów nie wymaga budowania strony od zera. "
    "To korekty — techniczne, architektoniczne, treściowe. "
    "Pierwsze efekty (lepsze pozycje) są widoczne w 60–90 dni.",
    bg=C_LGREEN, bar=C_GREEN, bold=False
)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# CZĘŚĆ 3: SZANSE
# CRO rule: "pieniądze leżące na stole" — konkretne, z liczbami
# Porządek: najszybszy efekt → długoterminowy
# ─────────────────────────────────────────────────────

H1(doc,"3. Gdzie możesz urosnąć?")
body(doc,
    "Trzy konkretne szanse — posegregowane od najszybszego efektu do długoterminowego "
    "budowania pozycji.",
    before=4,after=10
)

opportunity_card(doc,"A","60–90 dni",
    "57 fraz tuż za pierwszą stroną Google",
    "Masz 57 fraz, które siedzą na pozycjach 11–20 (tuż za pierwszą stroną) "
    "z niską konkurencją. Optymalizacja istniejących stron — bez nowych artykułów — "
    "może przenieść je do Top 10.",
    potential="Awans do Top 10 bez pisania nowych treści"
)

opportunity_card(doc,"B","3–6 miesięcy",
    "Tematy, które mają Twoi konkurenci, a Ty nie",
    "Megaland.pl, antonilacki.com i eventypilkarskie.pl rankują na frazy firmowe "
    "(\"atrakcje na imprezy firmowe Warszawa\", \"cennik dmuchańców\") "
    "o 3–10× wyższym wolumenie niż Twoje frazy dmuchańcowe. "
    "To nieexploatowany rynek B2B.",
    potential="500–3 000 wyszukiwań/mies. na jedną frazę — segement B2B z wyższym budżetem"
)

opportunity_card(doc,"C","1–2 miesiące",
    "Pytania, na które Google czeka — a Ty nie odpowiadasz",
    "Google pokazuje nam dokładnie co wpisują Twoi klienci: "
    "\"ile kosztuje wynajem dmuchańca?\", \"czy trzeba pozwolenie?\". "
    "Żadna z tych odpowiedzi nie istnieje na spogle.pl. "
    "Jedna strona FAQ z schema markup = szansa na Featured Snippet.",
    potential="Wyróżniony snippet i AI Overview — żaden z konkurentów go nie ma"
)

gap(doc,6)

# Quick wins tabela — konkretne frazy z liczbami
H2(doc,"Frazy gotowe do awansu — Top 5 Quick Wins")
body(doc,
    "Poniżej 5 najlepszych przykładów z 57 zidentyfikowanych szans. "
    "\"Trudność\" mierzy jak trudno wejść na 1. stronę Google (0–100, niżej = łatwiej):",
    after=6
)
table(doc,
    headers=["Fraza","Poz. teraz","Wyszukiwań/mies.","Trudność","Działanie"],
    rows=[
        ["wynajem dmuchańców na urodziny","16","320","31 ★ łatwa","Konsolidacja 6 stron lokalnych"],
        ["wypożyczalnia dmuchanych zjeżdżalni","15","170","33 ★ łatwa","Canonical + rozbudowa"],
        ["dmuchańce łódź","11","260","30 ★ łatwa","Dedykowana podstrona"],
        ["majówka z dziećmi","19","390","34 ★ łatwa","Aktualizacja + FAQ"],
        ["fajne zadania do wykonania","20","210","33 ★ łatwa","Rozbudowa + CTA do oferty"],
    ],
    widths=[5.5,2.8,4,3,4.7],
)

body(doc,"★ Trudność poniżej 35 = szanse bez dużego budżetu na linki zewnętrzne",
    color=C_TEAL,sz=9.5,before=0,after=6)

# Luki tematyczne B2B
H2(doc,"Tematy B2B, które mają Twoi konkurenci")
table(doc,
    headers=["Temat","Kto rankuje?","Wyszukiwań/mies.","Priorytet"],
    rows=[
        ["Cennik wynajmu dmuchańców 2026","megaland.pl, danmel.com.pl","1 000–3 000","🔴 P1"],
        ["Atrakcje na imprezy firmowe Warszawa","eventypilkarskie.pl, hulakula.com.pl","500–1 500","🔴 P1"],
        ["Wynajem dmuchańców Warszawa — rezerwacja","danmel.com.pl","800–2 000","🔴 P1"],
        ["Team building gry integracyjne","antonilacki.com","400–1 000","🟠 P2"],
        ["Bubble football — cena, zasady","eventypilkarskie.pl","300–700","🟠 P2"],
    ],
    widths=[5.5,4.5,3.5,3.5],
)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# CZĘŚĆ 4: PLAN DZIAŁAŃ
# CRO rule: konkretny, z timeframes, action-oriented copy
# Każde działanie = problem który rozwiązuje + spodziewany efekt
# ─────────────────────────────────────────────────────

H1(doc,"4. Plan działań")
body(doc,
    "Konkretne kroki posegregowane według priorytetu. "
    "P1 = robimy pierwsze — największy efekt przy najmniejszym nakładzie.",
    before=4,after=10
)

action_row(doc,"P1","30 dni",C_RED,
    "Naprawa kanibalizacji 59 fraz",
    "Przeprowadzamy audyt stron lokalnych i wskazujemy które połączyć lub przekierować (301). "
    "Zero pisania nowych treści. Efekt: \"wynajem dmuchańców na urodziny\" (poz. 16, 320/mies.) "
    "wchodzi do Top 10 w 60–90 dni."
)
action_row(doc,"P1","30 dni",C_RED,
    "Strona /cennik-wynajmu-dmuchancow + FAQ",
    "Brief + treść strony z cenami i odpowiedziami na pytania klientów. "
    "Schema FAQPage = szansa na Featured Snippet i wyświetlanie w Google AI Overview. "
    "Żaden z bezpośrednich konkurentów tej strony nie ma."
)
action_row(doc,"P1","30 dni",C_RED,
    "Przebudowa artykułu COVID 2020",
    "Artykuł rankuje na 109 fraz bez wartości komercyjnej. "
    "Zmiana tematu na \"8 gier integracyjnych na imprezy firmowe\" — "
    "zachowujemy pozycje, zamieniamy ruch informacyjny na biznesowy."
)

gap(doc,4)
action_row(doc,"P2","30–60 dni",C_ORANGE,
    "E-E-A-T: sygnały ekspertyzy w top 3 artykułach",
    "Imiona autorów, biogramy z 10-letnim doświadczeniem, liczba eventów 1000+, "
    "zdjęcia z realizacji, daty aktualizacji. Ochrona przed kolejnymi Core Updates Google."
)
action_row(doc,"P2","30–60 dni",C_ORANGE,
    "Landing page: Atrakcje na imprezy firmowe Warszawa",
    "Dedykowana strona atakująca frazy B2B (500–1 500 wyszukiwań/mies.). "
    "Z wyraźnym CTA do rezerwacji. Segment firmowy = wyższy budżet klienta."
)

gap(doc,4)
action_row(doc,"P3","60–90 dni",C_YELLOW,
    "Stworzenie strony /o-nas",
    "Aktualnie 404. Spogle, Maciek, Hubert, 10 lat, 1000+ eventów, Warszawa + Mazowsze. "
    "Kluczowy sygnał E-E-A-T — Google i klienci porównujący oferty go szukają."
)
action_row(doc,"P3","60–90 dni",C_YELLOW,
    "Optymalizacja 5 quick wins + rozbudowa stron produktowych",
    "Dla 5 najlepszych fraz z poz. 11–20: aktualizacja tytułu, H1, uzupełnienie treści. "
    "30–60 min per strona. Efekt w 4–8 tygodniach od wdrożenia."
)

gap(doc,4)
action_row(doc,"P4","3–12 mies.",C_TEAL,
    "Systematyczny content B2B + link building",
    "Miesięczny cykl artykułów: team building, imprezy firmowe, integracja pracowników. "
    "Równolegle: linki z portali branżowych. Cel: Domain Rank 52 387 → 35 000 w 12 mies."
)

gap(doc,8)

# Timeline — CRO: "kiedy zobaczę efekty?" to najczęstsza obiekcja
H2(doc,"Kiedy zobaczysz wyniki?")
table(doc,
    headers=["Działanie","Start","Pierwsze efekty","Cel"],
    rows=[
        ["Naprawa kanibalizacji","Tydzień 1–2","60–90 dni","Frazy komercyjne w Top 10"],
        ["Strona cennikowa + FAQ","Tydzień 2–3","2–3 miesiące","Featured Snippet, konwersja"],
        ["Landing B2B + E-E-A-T","Miesiąc 2","3–5 miesięcy","Ruch firmowy, ochrona pozycji"],
        ["Content B2B + link building","Miesiąc 3+","6–12 miesięcy","Nowe frazy Top 10, autorytet"],
        ["Powrót do 600+ fraz w Top 10","—","6–12 miesięcy","Cel roczny"],
    ],
    widths=[5.5,3,3.5,5],
)
body(doc,
    "Terminy szacunkowe — zależą od szybkości wdrożeń po stronie klienta.",
    color=C_GREY,sz=9,italic=True
)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# CZĘŚĆ 5: OBIEKCJE + CTA
# CRO rule: address objections BEFORE the ask — potem jasny następny krok
# ─────────────────────────────────────────────────────

H1(doc,"5. Pytania, które zazwyczaj pojawiają się na tym etapie")
body(doc,
    "Odpowiadamy na najczęstsze wątpliwości, zanim padną.",
    before=4,after=10
)

objection_row(doc,
    "A co jeśli Google znowu zaktualizuje algorytm?",
    "To właśnie jest powód, żeby działać — nie czekać. "
    "Strony z E-E-A-T, dobrą strukturą i realną ekspertyzą są odporne na Core Updates. "
    "Twój spadek w maju 2025 to sygnał że ta odporność wymaga wzmocnienia."
)
objection_row(doc,
    "Ile czasu zajmie mi wdrożenie tych zmian?",
    "Działania P1 (kanibalizacja + FAQ) wymagają 1–2 tygodnie pracy technicznej. "
    "Resztę dostarczamy w postaci gotowych briefów i materiałów do wdrożenia. "
    "Nasz model: minimum angażowania Twojego zespołu."
)
objection_row(doc,
    "Skąd wiem, że to zadziała?",
    "Wszystkie rekomendacje opieramy na danych z Senuto (nie opiniach). "
    "57 quick wins istnieje w bazie danych — to nie prognoza, to identyfikacja fraz "
    "które już rankujesz, tylko za nisko. Awans to kwestia optymalizacji, nie magii."
)
objection_row(doc,
    "Co jeśli konkurencja mnie wyprzedzi podczas gdy będziemy pracować?",
    "Megaland.pl rośnie +58% r/r. Każdy miesiąc bez działań to miesiąc "
    "w którym skracają dystans. Działania P1 są najtańszą obroną — "
    "naprawiają to co już istnieje, zanim konkurent to zajmie."
)

gap(doc,10)

# CTA — CRO rule: jeden jasny następny krok, benefit-focused, nie feature-focused
p=doc.add_paragraph(); sp(p,before=0,after=0)
shd_para(p._p,C_DARK)
p.paragraph_format.left_indent=Cm(0.6)
r(p,"  Następny krok",bold=True,sz=9,color=C_ACCENT)

p=doc.add_paragraph(); sp(p,before=0,after=0)
shd_para(p._p,C_DARK)
p.paragraph_format.left_indent=Cm(0.6)
r(p,"  Umów 30-minutowe spotkanie — omówimy priorytety P1",
    bold=True,sz=15,color=C_WHITE)

p=doc.add_paragraph(); sp(p,before=0,after=0)
shd_para(p._p,C_DARK)
p.paragraph_format.left_indent=Cm(0.6)
r(p,"  i pokażemy demo naprawy kanibalizacji dla spogle.pl.",
    sz=12,color=RGBColor(0xCC,0xCC,0xCC))

p=doc.add_paragraph(); sp(p,before=0,after=6)
shd_para(p._p,C_DARK)
p.paragraph_format.left_indent=Cm(0.6)
r(p,"  double-digital.pl  ·  kontakt@double-digital.pl",
    sz=10,color=C_ACCENT,italic=True)

gap(doc,10)

# Social proof DD — CRO: trust signals blisko CTA
callout(doc,
    "Double Digital — Google Partner, 25+ krajów, od 2022. "
    "Specjalizacja: data-driven SEO, performance marketing, GA4 / BigQuery. "
    "Filozofia: \"Double or Nothing\" — konkretne wyniki, mierzalne KPI, "
    "miesięczna transparentność (raport Senuto + spotkanie statusowe).",
    bg=C_LGREY, bar=C_ACCENT
)

doc.add_page_break()

# ─────────────────────────────────────────────────────
# CZĘŚĆ 6: DANE SZCZEGÓŁOWE (appendix)
# CRO rule: dane dla "researcher" — nie na początku, ale dostępne
# Klient który chce zweryfikować — znajdzie tu wszystko
# ─────────────────────────────────────────────────────

H1(doc,"6. Dane źródłowe — dla zainteresowanych")
body(doc,
    "Ta sekcja zawiera surowe liczby z Senuto, na których oparliśmy całą analizę. "
    "Dla tych, którzy chcą zweryfikować każdą rekomendację.",
    before=4,after=8
)

H3(doc,"Senuto — snapshot widoczności (27.02.2026)")
table(doc,
    headers=["Metryka","Wartość","Poprzedni tydzień","Zmiana"],
    rows=[
        ["Widoczność","3 677 pkt","3 679 pkt","–0,05%"],
        ["Frazy w Top 3","207","207","0%"],
        ["Frazy w Top 10","523","524","–0,19%"],
        ["Frazy w Top 50","1 943","1 946","–0,15%"],
        ["Domain Rank","52 387","52 242","+0,28% (gorszy)"],
        ["Ads Equivalent","4 755 PLN/mies.","4 759 PLN/mies.","–0,08%"],
    ],
    widths=[5,4,4,3],
)

H3(doc,"Profil trudności słów kluczowych")
table(doc,
    headers=["Segment KD","Frazy (próba 500)","% portfolio","Znaczenie dla strategii"],
    rows=[
        ["Łatwe (KD < 30)","~155","31%","Quick wins — szybkie efekty optymalizacji"],
        ["Średnie (KD 30–60)","~340","68%","Wymagają treści + linków"],
        ["Trudne (KD > 60)","~5","1%","Pomijalne na tym etapie"],
    ],
    widths=[4,4,3,6],
)

H3(doc,"Pełna lista konkurentów (Senuto)")
table(doc,
    headers=["Domena","Wspólne frazy","Widoczność","Top 10","Domain Rank","Trend r/r"],
    rows=[
        ["megaland.pl","4","844","84","145 588","▲ +58,6% ⚠"],
        ["antonilacki.com","3","739","36","158 002","→ stabilny"],
        ["eventypilkarskie.pl","4","382","103","233 817","→ stabilny"],
        ["danmel.com.pl","3","223","47","313 245","→ stabilny"],
    ],
    widths=[4.5,3.5,3.5,2.5,3.5,3.5],
)

H3(doc,"Metodologia")
for m in [
    "Dane SEO: Senuto Visibility Analysis, Base 2.0, Polska (country_id=200) — 27.02.2026",
    "Dane SERP: NodeHub API, Google PL — 27.02.2026",
    "Analizowane frazy SERP: \"wynajem dmuchańców na urodziny\", \"wypożyczalnia dmuchanych zjeżdżalni\", \"atrakcje na imprezy firmowe Warszawa\"",
    "Analiza treści: Jina Reader (scraping) + ocena LLM",
    "Quick wins: pozycje 11–20 + KD < 35; kanibalizacja: HIGH = >20 fraz; próba KD: 500 z 2 137 fraz",
]:
    bul(doc,m)

gap(doc,10)
p=doc.add_paragraph(); sp(p,before=0,after=2)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r(p,"Double Digital  ·  double-digital.pl",sz=9,color=C_GREY,italic=True)

# SAVE
out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"Oferta_SEO_Spogle_DoubleDigital.docx")
doc.save(out)
print(f"✅ Zapisano: {out}")
